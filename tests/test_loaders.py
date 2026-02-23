import pytest
import io
import docx
from unittest.mock import patch
from pypdf import PdfWriter
from langchain_core.documents import Document
from ingestion.loaders import (
    load_urls,
    load_pdf,
    load_docx,
    load_txt,
    load_raw_text,
    load_all_sources,
)


# --- Mocking Streamlit UploadedFile ---
class DummyUploadedFile(io.BytesIO):
    """Mocks the Streamlit UploadedFile object behavior correctly as a file stream."""

    def __init__(self, name, content_bytes):
        super().__init__(content_bytes)
        self.name = name


# --- Fixtures for Dynamic File Generation ---
@pytest.fixture
def sample_txt_file():
    content = b"This is a valid text file content."
    return DummyUploadedFile("test.txt", content)


@pytest.fixture
def sample_docx_file():
    doc = docx.Document()
    doc.add_paragraph("This is a valid word document.")
    byte_stream = io.BytesIO()
    doc.save(byte_stream)
    return DummyUploadedFile("test.docx", byte_stream.getvalue())


@pytest.fixture
def sample_pdf_file():
    writer = PdfWriter()
    writer.add_blank_page(width=72, height=72)
    byte_stream = io.BytesIO()
    writer.write(byte_stream)
    return DummyUploadedFile("test.pdf", byte_stream.getvalue())


@pytest.fixture
def invalid_file():
    content = b"\x00\x01\x02\x03corrupted_bytes"
    return DummyUploadedFile("corrupt.file", content)


# --- Tests for Raw Text Handler ---
def test_load_raw_text_valid():
    text = "Here is some raw input."
    docs = load_raw_text(text)
    assert len(docs) == 1
    assert isinstance(docs[0], Document)
    assert docs[0].page_content == "Here is some raw input."
    assert docs[0].metadata["source_type"] == "raw_text"


def test_load_raw_text_empty():
    docs = load_raw_text("   ")
    assert len(docs) == 0


# --- Tests for TXT Loader ---
def test_load_txt_valid(sample_txt_file):
    docs = load_txt(sample_txt_file)
    assert len(docs) == 1
    assert "valid text file content" in docs[0].page_content
    assert docs[0].metadata["source_type"] == "txt"


def test_load_txt_empty():
    docs = load_txt(None)
    assert len(docs) == 0


# --- Tests for DOCX Loader ---
def test_load_docx_valid(sample_docx_file):
    docs = load_docx(sample_docx_file)
    assert len(docs) == 1
    assert "valid word document" in docs[0].page_content
    assert docs[0].metadata["source_type"] == "docx"


def test_load_docx_invalid(invalid_file):
    with pytest.raises(RuntimeError):
        load_docx(invalid_file)


# --- Tests for PDF Loader ---
def test_load_pdf_valid(sample_pdf_file):
    docs = load_pdf(sample_pdf_file)
    assert len(docs) == 1
    assert docs[0].metadata["source_type"] == "pdf"
    assert "page" in docs[0].metadata


def test_load_pdf_invalid(invalid_file):
    with pytest.raises(RuntimeError):
        load_pdf(invalid_file)


# --- Tests for URL Loader ---
@patch("ingestion.loaders.WebBaseLoader.load")
def test_load_urls_valid(mock_load):
    # Mock the network call to avoid SSL or internet dependency issues in tests
    mock_load.return_value = [
        Document(
            page_content="Mock web content", metadata={"source": "https://example.com"}
        )
    ]
    urls = ["https://example.com"]
    docs = load_urls(urls)
    assert len(docs) > 0
    assert docs[0].metadata["source_type"] == "url"
    assert docs[0].metadata["source_name"] == "https://example.com"


def test_load_urls_invalid():
    # Loaders are designed to warn and skip invalid URLs, returning empty list
    urls = ["https://this-is-a-completely-invalid-url-that-does-not-exist.com"]
    docs = load_urls(urls)
    assert len(docs) == 0


# --- Tests for Master Orchestrator ---
def test_load_all_sources_empty():
    with pytest.raises(
        ValueError, match="No valid document content could be extracted"
    ):
        load_all_sources()


def test_load_all_sources_combined(sample_txt_file, sample_docx_file):
    docs = load_all_sources(
        txt_files=[sample_txt_file],
        docx_files=[sample_docx_file],
        raw_text="Additional raw text",
    )
    # Should extract 1 from TXT, 1 from DOCX, 1 from Raw Text
    assert len(docs) == 3
