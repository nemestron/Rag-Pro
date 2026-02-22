import os
import tempfile
import logging
from typing import List, Dict, Any, Optional

from langchain_core.documents import Document
from langchain_community.document_loaders import WebBaseLoader, PyPDFLoader
import docx

# Configure module-level logger
logger = logging.getLogger(__name__)

def load_urls(urls: List[str]) -> List[Document]:
    """
    Load content from a list of URLs using WebBaseLoader.
    """
    documents = []
    if not urls:
        return documents

    for url in urls:
        url = url.strip()
        if not url:
            continue
            
        try:
            logger.info(f"Loading URL: {url}")
            loader = WebBaseLoader(web_paths=[url])
            docs = loader.load()
            
            # Ensure metadata adheres to the contract
            for doc in docs:
                doc.metadata["source_type"] = "url"
                doc.metadata["source_name"] = url
                
            documents.extend(docs)
        except Exception as e:
            logger.warning(f"Failed to load URL {url}: {str(e)}")
            
    return documents

def load_pdf(uploaded_file) -> List[Document]:
    """
    Load content from a Streamlit UploadedFile object representing a PDF.
    """
    documents = []
    if not uploaded_file:
        return documents

    file_name = uploaded_file.name
    logger.info(f"Loading PDF: {file_name}")
    
    # PyPDFLoader requires a file path, so we write the in-memory file to a temp file
    try:
        with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as temp_file:
            temp_file.write(uploaded_file.getvalue())
            temp_path = temp_file.name

        try:
            loader = PyPDFLoader(file_path=temp_path)
            docs = loader.load()
            
            for doc in docs:
                doc.metadata["source_type"] = "pdf"
                doc.metadata["source_name"] = file_name
                # PyPDFLoader automatically adds 'page' to metadata
                
            documents.extend(docs)
        finally:
            # Ensure temporary file is cleaned up regardless of extraction success
            if os.path.exists(temp_path):
                os.remove(temp_path)
                
    except Exception as e:
        logger.error(f"Failed to process PDF {file_name}: {str(e)}")
        
    return documents

def load_docx(uploaded_file) -> List[Document]:
    """
    Load content from a Streamlit UploadedFile object representing a DOCX.
    """
    documents = []
    if not uploaded_file:
        return documents

    file_name = uploaded_file.name
    logger.info(f"Loading DOCX: {file_name}")
    
    try:
        # python-docx can read directly from a file-like object
        doc = docx.Document(uploaded_file)
        full_text = []
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text)
                
        content = "\n".join(full_text)
        
        if content.strip():
            metadata = {
                "source_type": "docx",
                "source_name": file_name
            }
            documents.append(Document(page_content=content, metadata=metadata))
            
    except Exception as e:
        logger.error(f"Failed to process DOCX {file_name}: {str(e)}")
        
    return documents

def load_txt(uploaded_file) -> List[Document]:
    """
    Load content from a Streamlit UploadedFile object representing a TXT file.
    """
    documents = []
    if not uploaded_file:
        return documents

    file_name = uploaded_file.name
    logger.info(f"Loading TXT: {file_name}")
    
    try:
        bytes_data = uploaded_file.getvalue()
        # Attempt UTF-8 decoding first, fallback to latin-1
        try:
            content = bytes_data.decode("utf-8")
        except UnicodeDecodeError:
            content = bytes_data.decode("latin-1")
            
        if content.strip():
            metadata = {
                "source_type": "txt",
                "source_name": file_name
            }
            documents.append(Document(page_content=content, metadata=metadata))
            
    except Exception as e:
        logger.error(f"Failed to process TXT {file_name}: {str(e)}")
        
    return documents

def load_raw_text(text: str) -> List[Document]:
    """
    Wrap direct user text input into a Document object.
    """
    if not text or not text.strip():
        return []
        
    logger.info("Loading raw text input")
    metadata = {
        "source_type": "raw_text",
        "source_name": "user_input"
    }
    return [Document(page_content=text.strip(), metadata=metadata)]

def load_all_sources(
    urls: Optional[List[str]] = None,
    pdf_files: Optional[List[Any]] = None,
    docx_files: Optional[List[Any]] = None,
    txt_files: Optional[List[Any]] = None,
    raw_text: Optional[str] = ""
) -> List[Document]:
    """
    Master orchestrator function to load all provided sources into a unified list of Documents.
    """
    all_documents: List[Document] = []
    
    # Process URLs
    if urls:
        all_documents.extend(load_urls(urls))
        
    # Process PDFs
    if pdf_files:
        for pdf in pdf_files:
            all_documents.extend(load_pdf(pdf))
            
    # Process DOCX files
    if docx_files:
        for docx_file in docx_files:
            all_documents.extend(load_docx(docx_file))
            
    # Process TXT files
    if txt_files:
        for txt_file in txt_files:
            all_documents.extend(load_txt(txt_file))
            
    # Process Raw Text
    if raw_text:
        all_documents.extend(load_raw_text(raw_text))
        
    if not all_documents:
        raise ValueError("No valid document content could be extracted from the provided sources.")
        
    logger.info(f"Successfully loaded a total of {len(all_documents)} document objects across all sources.")
    return all_documents
